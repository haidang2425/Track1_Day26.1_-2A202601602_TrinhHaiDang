import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn
import subprocess
import os

def set_cell_background(cell, fill_hex):
    tcPr = cell._element.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=60, bottom=60, left=100, right=100):
    tcPr = cell._element.get_or_add_tcPr()
    tcMar = parse_xml(
        f'<w:tcMar {nsdecls("w")}>'
        f'<w:top w:w="{top}" w:type="dxa"/>'
        f'<w:bottom w:w="{bottom}" w:type="dxa"/>'
        f'<w:left w:w="{left}" w:type="dxa"/>'
        f'<w:right w:w="{right}" w:type="dxa"/>'
        f'</w:tcMar>'
    )
    tcPr.append(tcMar)

def set_table_borders(table, color="D1D5DB", sz="4", val="single"):
    tblPr = table._element.xpath('w:tblPr')
    if tblPr:
        borders = parse_xml(
            f'<w:tblBorders {nsdecls("w")}>'
            f'<w:top w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>'
            f'<w:bottom w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>'
            f'<w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
            f'<w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
            f'<w:insideH w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>'
            f'<w:insideV w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>'
            f'</w:tblBorders>'
        )
        tblPr[0].append(borders)

def format_para(para, font_name="Arial", size_pt=7.5, bold=False, italic=False, color_rgb=RGBColor(31, 41, 55), align=WD_ALIGN_PARAGRAPH.LEFT, space_before=0, space_after=0, line_spacing=1.05):
    para.alignment = align
    para.paragraph_format.space_before = Pt(space_before)
    para.paragraph_format.space_after = Pt(space_after)
    para.paragraph_format.line_spacing = line_spacing
    for run in para.runs:
        run.font.name = font_name
        run.font.size = Pt(size_pt)
        run.font.bold = bold
        run.font.italic = italic
        run.font.color.rgb = color_rgb

def add_styled_cell(cell, text_runs, fill_hex=None, align=WD_ALIGN_PARAGRAPH.LEFT, top_pad=40, bot_pad=40):
    set_cell_margins(cell, top=top_pad, bottom=bot_pad, left=80, right=80)
    if fill_hex:
        set_cell_background(cell, fill_hex)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    para = cell.paragraphs[0]
    para.alignment = align
    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.space_after = Pt(0)
    para.paragraph_format.line_spacing = 1.05
    
    for text, bold, italic, size, color in text_runs:
        run = para.add_run(text)
        run.font.name = "Arial"
        run.font.size = Pt(size)
        run.font.bold = bold
        run.font.italic = italic
        run.font.color.rgb = color

def create_document():
    doc = docx.Document()
    
    # Set standard A4 with compact executive margins
    section = doc.sections[0]
    section.page_width = Inches(8.27)
    section.page_height = Inches(11.69)
    section.top_margin = Inches(0.35)
    section.bottom_margin = Inches(0.35)
    section.left_margin = Inches(0.4)
    section.right_margin = Inches(0.4)
    
    # ------------------ PAGE 1: OPERATING DASHBOARD ------------------
    # Main Header
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p_title.paragraph_format.space_before = Pt(0)
    p_title.paragraph_format.space_after = Pt(1)
    r = p_title.add_run("BẢNG ĐIỀU KHIỂN VẬN HÀNH (OPERATING DASHBOARD) — CURSUS AI")
    r.font.name = "Arial"
    r.font.size = Pt(11.5)
    r.font.bold = True
    r.font.color.rgb = RGBColor(15, 41, 74)

    # Metadata Line
    p_meta = doc.add_paragraph()
    p_meta.paragraph_format.space_before = Pt(0)
    p_meta.paragraph_format.space_after = Pt(3)
    p_meta.paragraph_format.line_spacing = 1.1
    runs_meta = [
        ("Học viên: ", True, False, 7.5, RGBColor(15, 41, 74)),
        ("Trịnh Hải Đăng  |  ", False, False, 7.5, RGBColor(31, 41, 55)),
        ("Mã học viên: ", True, False, 7.5, RGBColor(15, 41, 74)),
        ("2A202601602  |  ", False, False, 7.5, RGBColor(31, 41, 55)),
        ("Mô hình: ", True, False, 7.5, RGBColor(15, 41, 74)),
        ("B2B (Higher Education)  |  ", False, False, 7.5, RGBColor(31, 41, 55)),
        ("Cập nhật: ", True, False, 7.5, RGBColor(15, 41, 74)),
        ("2026-08-28  |  ", False, False, 7.5, RGBColor(31, 41, 55)),
        ("Owner: ", True, False, 7.5, RGBColor(15, 41, 74)),
        ("Product Operations Lead", False, False, 7.5, RGBColor(31, 41, 55)),
    ]
    for text, b, it, sz, clr in runs_meta:
        r = p_meta.add_run(text)
        r.font.name = "Arial"
        r.font.size = Pt(sz)
        r.font.bold = b
        r.font.italic = it
        r.font.color.rgb = clr

    # Diagnosis & North Star Callout Table
    tbl_diag = doc.add_table(rows=1, cols=1)
    tbl_diag.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl_diag.autofit = False
    tbl_diag.columns[0].width = Inches(7.47)
    cell_diag = tbl_diag.cell(0, 0)
    set_cell_margins(cell_diag, top=50, bottom=50, left=90, right=90)
    set_cell_background(cell_diag, "F1F5F9")
    p_d = cell_diag.paragraphs[0]
    p_d.paragraph_format.space_before = Pt(0)
    p_d.paragraph_format.space_after = Pt(0)
    p_d.paragraph_format.line_spacing = 1.15
    diag_content = [
        ("Chẩn đoán mô hình: ", True, False, 7.2, RGBColor(15, 41, 74)),
        ("B2B do Khoa/Trường Đại học ký hợp đồng & chi trả định kỳ ($14.400 ACV/khoa) từ ngân sách vận hành đào tạo & quỹ TA; Giảng viên và Sinh viên sử dụng trực tiếp qua chuẩn Canvas LMS LTI 1.3.\n", False, False, 7.2, RGBColor(31, 41, 55)),
        ("North Star Metric: ", True, False, 7.2, RGBColor(15, 41, 74)),
        ("Time-to-first-value ≤ 5 ngày với ít nhất 50 sinh viên active trong khóa học (Hiện tại: ", False, False, 7.2, RGBColor(31, 41, 55)),
        ("6 ngày", True, False, 7.2, RGBColor(180, 83, 9)),
        (" — Trạng thái: Cảnh báo)", False, False, 7.2, RGBColor(31, 41, 55))
    ]
    for text, b, it, sz, clr in diag_content:
        r = p_d.add_run(text)
        r.font.name = "Arial"
        r.font.size = Pt(sz)
        r.font.bold = b
        r.font.italic = it
        r.font.color.rgb = clr

    # Spacer
    p_sp = doc.add_paragraph()
    p_sp.paragraph_format.space_before = Pt(2)
    p_sp.paragraph_format.space_after = Pt(1)

    # 1. Section 1 Heading
    p_h1 = doc.add_paragraph()
    p_h1.paragraph_format.space_before = Pt(0)
    p_h1.paragraph_format.space_after = Pt(2)
    r = p_h1.add_run("1. CÂY TÍN HIỆU 3 TẦNG & CÁC NGƯỠNG VẬN HÀNH (Signal Tree & Thresholds)")
    r.font.name = "Arial"
    r.font.size = Pt(8.5)
    r.font.bold = True
    r.font.color.rgb = RGBColor(15, 41, 74)

    # Table 1: 3-Tier Dashboard
    # Columns: [Tầng, ID, Tên chỉ số & Định nghĩa, Hiện tại · Nguồn, Ngưỡng Xanh, Ngưỡng Vàng, Ngưỡng Đỏ, Nhịp · Owner, Báo trước cho · Luật]
    col_widths_t1 = [Inches(0.40), Inches(0.40), Inches(1.85), Inches(0.95), Inches(0.70), Inches(0.75), Inches(0.65), Inches(0.85), Inches(0.92)]
    tbl1 = doc.add_table(rows=8, cols=9)
    tbl1.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl1.autofit = False
    set_table_borders(tbl1, color="CBD5E1", sz="4")

    headers_t1 = ["Tầng", "ID", "Tên chỉ số & Định nghĩa", "Hiện tại · Nguồn", "Vùng Xanh", "Vùng Vàng", "Vùng Đỏ", "Nhịp · Owner", "Báo trước · Luật"]
    for j, h in enumerate(headers_t1):
        cell = tbl1.cell(0, j)
        cell.width = col_widths_t1[j]
        add_styled_cell(cell, [(h, True, False, 6.8, RGBColor(255, 255, 255))], fill_hex="0F294A", top_pad=40, bot_pad=40)

    data_t1 = [
        ("Leading", "L-01", [("Time-to-first-value (TTFV)\n", True, False, 6.7, RGBColor(17, 24, 39)), ("Số ngày từ kết nối LTI đến khi đạt 50 resolution QA", False, False, 6.4, RGBColor(75, 85, 99))], "6 ngày\n[TB] Pilot cohort", "≤ 5 ngày", "6–10 ngày", "> 10 ngày", "Tuần\nProduct Ops", "POC-to-paid & SV\nR-01 (Dừng)"),
        ("Leading", "L-02", [("Weekly student resolution rate\n", True, False, 6.7, RGBColor(17, 24, 39)), ("Tỷ lệ SV active có ≥2 resolution QA / tuần", False, False, 6.4, RGBColor(75, 85, 99))], "58%\n[TB] ĐHBK log", "≥ 65%", "45–64%", "< 45%", "Tuần\nCustomer Success", "TTFV & Churn\nR-02"),
        ("Operating", "O-01", [("Pilot activation rate\n", True, False, 6.7, RGBColor(17, 24, 39)), ("Tỷ lệ pilot khoa đạt >500 resolution trong 30 ngày", False, False, 6.4, RGBColor(75, 85, 99))], "66%\n[MH] MH-02", "≥ 75%", "50–74%", "< 50%", "Tuần\nProduct Ops", "POC-to-paid\nR-03 (Dừng)"),
        ("Operating", "O-02", [("Chi phí AI trên mỗi job\n", True, False, 6.7, RGBColor(17, 24, 39)), ("Tổng token Claude Haiku + Infra chia resolution QA", False, False, 6.4, RGBColor(75, 85, 99))], "840 đ\n[MH] MH-01", "≤ 1.000 đ", "1.001–1.800 đ", "> 1.800 đ", "Tuần\nFinOps", "Gross Margin\nR-04"),
        ("Operating", "O-03", [("POC-to-paid conversion\n", True, False, 6.7, RGBColor(17, 24, 39)), ("Tỷ lệ pilot chuyển đổi thành hợp đồng năm chính thức", False, False, 6.4, RGBColor(75, 85, 99))], "50%\n[BM] ICONIQ '26", "≥ 60%", "40–59%", "< 40%", "Tháng\nRevenue Ops", "ARR & CAC payback\nR-03 (Dừng)"),
        ("Lagging", "G-01", [("Gross margin sau chi phí AI\n", True, False, 6.7, RGBColor(17, 24, 39)), ("(Doanh thu − toàn bộ COGS AI) ÷ Doanh thu", False, False, 6.4, RGBColor(75, 85, 99))], "82%\n[MH] MH-01", "≥ 80%", "65–79%", "< 65%", "Tháng\nFinance", "Runway & Payback\nR-04"),
        ("Lagging", "G-02", [("Net Revenue Retention (NRR)\n", True, False, 6.7, RGBColor(17, 24, 39)), ("Doanh thu cohort sau mở rộng môn/lớp và churn", False, False, 6.4, RGBColor(75, 85, 99))], "100%\n[BM] Benchmarkit", "≥ 110%", "100–109%", "< 100%", "Quý\nFinance", "LTV & Tăng trưởng\nR-05")
    ]

    for i, row in enumerate(data_t1, start=1):
        bg = "FFFFFF" if i % 2 != 0 else "F8FAFC"
        tbl1.cell(i, 0).width = col_widths_t1[0]
        add_styled_cell(tbl1.cell(i, 0), [(row[0], False, False, 6.5, RGBColor(75, 85, 99))], fill_hex=bg)
        
        tbl1.cell(i, 1).width = col_widths_t1[1]
        add_styled_cell(tbl1.cell(i, 1), [(row[1], True, False, 6.7, RGBColor(15, 41, 74))], fill_hex=bg)
        
        tbl1.cell(i, 2).width = col_widths_t1[2]
        add_styled_cell(tbl1.cell(i, 2), row[2], fill_hex=bg)
        
        tbl1.cell(i, 3).width = col_widths_t1[3]
        add_styled_cell(tbl1.cell(i, 3), [(row[3], False, False, 6.5, RGBColor(31, 41, 55))], fill_hex=bg)
        
        tbl1.cell(i, 4).width = col_widths_t1[4]
        add_styled_cell(tbl1.cell(i, 4), [(row[4], True, False, 6.6, RGBColor(4, 120, 87))], fill_hex=bg)
        
        tbl1.cell(i, 5).width = col_widths_t1[5]
        add_styled_cell(tbl1.cell(i, 5), [(row[5], False, False, 6.6, RGBColor(180, 83, 9))], fill_hex=bg)
        
        tbl1.cell(i, 6).width = col_widths_t1[6]
        add_styled_cell(tbl1.cell(i, 6), [(row[6], True, False, 6.6, RGBColor(185, 28, 28))], fill_hex=bg)
        
        tbl1.cell(i, 7).width = col_widths_t1[7]
        add_styled_cell(tbl1.cell(i, 7), [(row[7], False, False, 6.5, RGBColor(31, 41, 55))], fill_hex=bg)
        
        tbl1.cell(i, 8).width = col_widths_t1[8]
        add_styled_cell(tbl1.cell(i, 8), [(row[8], False, False, 6.5, RGBColor(31, 41, 55))], fill_hex=bg)

    # 2. Section 2 Heading
    p_h2 = doc.add_paragraph()
    p_h2.paragraph_format.space_before = Pt(3)
    p_h2.paragraph_format.space_after = Pt(2)
    r = p_h2.add_run("2. HỆ THỐNG LUẬT QUYẾT ĐỊNH (Decision Rules: NẾU − TRONG − VÀ − THÌ − KHÔNG THÌ)")
    r.font.name = "Arial"
    r.font.size = Pt(8.5)
    r.font.bold = True
    r.font.color.rgb = RGBColor(15, 41, 74)

    col_widths_t2 = [Inches(0.40), Inches(2.05), Inches(2.35), Inches(2.05), Inches(0.62)]
    tbl2 = doc.add_table(rows=6, cols=5)
    tbl2.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl2.autofit = False
    set_table_borders(tbl2, color="CBD5E1", sz="4")

    headers_t2 = ["Mã", "Điều kiện kích hoạt (NẾU · TRONG · VÀ)", "Hành động bắt buộc (THÌ)", "Phản xạ sai bị cấm (KHÔNG THÌ)", "Luật dừng"]
    for j, h in enumerate(headers_t2):
        cell = tbl2.cell(0, j)
        cell.width = col_widths_t2[j]
        add_styled_cell(cell, [(h, True, False, 6.8, RGBColor(255, 255, 255))], fill_hex="0F294A", top_pad=40, bot_pad=40)

    data_t2 = [
        ("R-01", [("NẾU ", True, False, 6.6, RGBColor(15, 41, 74)), ("Median TTFV > 10 ngày\n", True, False, 6.6, RGBColor(185, 28, 28)), ("TRONG: ", True, False, 6.4, RGBColor(75, 85, 99)), ("2 cohort liên tiếp\n", False, False, 6.4, RGBColor(31, 41, 55)), ("VÀ: ", True, False, 6.4, RGBColor(75, 85, 99)), ("Mỗi cohort có ≥2 khoa tham gia", False, False, 6.4, RGBColor(31, 41, 55))], [("Đóng băng tiếp nhận pilot mới trong 14 ngày", True, False, 6.6, RGBColor(17, 24, 39)), (" và tinh gọn quy trình cấu hình RAG syllabus xuống đúng 1 môn học cốt lõi", False, False, 6.4, RGBColor(31, 41, 55))], [("CẤM: ", True, False, 6.5, RGBColor(185, 28, 28)), ("Không giảm giá hợp đồng để bù đắp sự chậm trễ trong việc chứng minh giá trị", False, False, 6.4, RGBColor(31, 41, 55))], "CÓ (Dừng)"),
        ("R-02", [("NẾU ", True, False, 6.6, RGBColor(15, 41, 74)), ("Weekly resolution rate < 45%\n", True, False, 6.6, RGBColor(185, 28, 28)), ("TRONG: ", True, False, 6.4, RGBColor(75, 85, 99)), ("3 tuần liên tiếp\n", False, False, 6.4, RGBColor(31, 41, 55)), ("VÀ: ", True, False, 6.4, RGBColor(75, 85, 99)), ("Có ≥100 sinh viên trong danh sách lớp", False, False, 6.4, RGBColor(31, 41, 55))], [("Biệt phái 1 chuyên viên Product Ops", True, False, 6.6, RGBColor(17, 24, 39)), (" trực tiếp hỗ trợ giảng viên gắn bài tập tuần vào widget Canvas", False, False, 6.4, RGBColor(31, 41, 55))], [("CẤM: ", True, False, 6.5, RGBColor(185, 28, 28)), ("Không gửi email spam thúc ép sinh viên khi giao diện chưa thuận tiện", False, False, 6.4, RGBColor(31, 41, 55))], "KHÔNG"),
        ("R-03", [("NẾU ", True, False, 6.6, RGBColor(15, 41, 74)), ("Pilot activation rate < 50%\n", True, False, 6.6, RGBColor(185, 28, 28)), ("TRONG: ", True, False, 6.4, RGBColor(75, 85, 99)), ("2 kỳ đánh giá liên tiếp\n", False, False, 6.4, RGBColor(31, 41, 55)), ("VÀ: ", True, False, 6.4, RGBColor(75, 85, 99)), ("Có ≥4 pilot khoa đang chạy", False, False, 6.4, RGBColor(31, 41, 55))], [("Dừng toàn bộ hoạt động outbound sales mới", True, False, 6.6, RGBColor(17, 24, 39)), (" và tập trung tối ưu hóa kịch bản onboarding 1-click cho giảng viên", False, False, 6.4, RGBColor(31, 41, 55))], [("CẤM: ", True, False, 6.5, RGBColor(185, 28, 28)), ("Không tăng ngân sách sales để tìm thêm pilot khi tỷ lệ kích hoạt chưa đạt", False, False, 6.4, RGBColor(31, 41, 55))], "CÓ (Dừng)"),
        ("R-04", [("NẾU ", True, False, 6.6, RGBColor(15, 41, 74)), ("AI cost / job > 1.800 đ\n", True, False, 6.6, RGBColor(185, 28, 28)), ("TRONG: ", True, False, 6.4, RGBColor(75, 85, 99)), ("2 tuần liên tiếp\n", False, False, 6.4, RGBColor(31, 41, 55)), ("VÀ: ", True, False, 6.4, RGBColor(75, 85, 99)), ("Có ≥1.000 resolution phát sinh", False, False, 6.4, RGBColor(31, 41, 55))], [("Bật tính năng prompt caching", True, False, 6.6, RGBColor(17, 24, 39)), (", cắt giảm context RAG và chuyển truy vấn đơn giản sang model tier nhỏ", False, False, 6.4, RGBColor(31, 41, 55))], [("CẤM: ", True, False, 6.5, RGBColor(185, 28, 28)), ("Không tắt bộ lọc an toàn và kiểm duyệt chất lượng để giảm chi phí token ảo", False, False, 6.4, RGBColor(31, 41, 55))], "KHÔNG"),
        ("R-05", [("NẾU ", True, False, 6.6, RGBColor(15, 41, 74)), ("NRR < 100%\n", True, False, 6.6, RGBColor(185, 28, 28)), ("TRONG: ", True, False, 6.4, RGBColor(75, 85, 99)), ("2 quý liên tiếp\n", False, False, 6.4, RGBColor(31, 41, 55)), ("VÀ: ", True, False, 6.4, RGBColor(75, 85, 99)), ("Có ≥3 hợp đồng khoa đến kỳ gia hạn", False, False, 6.4, RGBColor(31, 41, 55))], [("Tổ chức phiên làm việc trực tiếp Trưởng khoa", True, False, 6.6, RGBColor(17, 24, 39)), (" đánh giá báo cáo giảm tải TA và giải quyết các rào cản tính năng", False, False, 6.4, RGBColor(31, 41, 55))], [("CẤM: ", True, False, 6.5, RGBColor(185, 28, 28)), ("Không tính toán gộp cơ hội bán mới từ các trường khác vào NRR tài khoản cũ", False, False, 6.4, RGBColor(31, 41, 55))], "KHÔNG")
    ]

    for i, row in enumerate(data_t2, start=1):
        bg = "FFFFFF" if i % 2 != 0 else "F8FAFC"
        tbl2.cell(i, 0).width = col_widths_t2[0]
        add_styled_cell(tbl2.cell(i, 0), [(row[0], True, False, 6.8, RGBColor(15, 41, 74))], fill_hex=bg)
        
        tbl2.cell(i, 1).width = col_widths_t2[1]
        add_styled_cell(tbl2.cell(i, 1), row[1], fill_hex=bg)
        
        tbl2.cell(i, 2).width = col_widths_t2[2]
        add_styled_cell(tbl2.cell(i, 2), row[2], fill_hex=bg)
        
        tbl2.cell(i, 3).width = col_widths_t2[3]
        add_styled_cell(tbl2.cell(i, 3), row[3], fill_hex=bg)
        
        tbl2.cell(i, 4).width = col_widths_t2[4]
        color_stop = RGBColor(185, 28, 28) if "CÓ" in row[4] else RGBColor(75, 85, 99)
        add_styled_cell(tbl2.cell(i, 4), [(row[4], True, False, 6.6, color_stop)], fill_hex=bg)

    # 3. Section 3 Heading
    p_h3 = doc.add_paragraph()
    p_h3.paragraph_format.space_before = Pt(3)
    p_h3.paragraph_format.space_after = Pt(2)
    r = p_h3.add_run("3. CỔNG GÁC 90 NGÀY, KILL CRITERIA & ĐIỂM CHƯA ĐO ĐƯỢC")
    r.font.name = "Arial"
    r.font.size = Pt(8.5)
    r.font.bold = True
    r.font.color.rgb = RGBColor(15, 41, 74)

    col_widths_t3 = [Inches(0.65), Inches(0.75), Inches(2.20), Inches(1.85), Inches(1.20), Inches(0.82)]
    tbl3 = doc.add_table(rows=4, cols=6)
    tbl3.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl3.autofit = False
    set_table_borders(tbl3, color="CBD5E1", sz="4")

    headers_t3 = ["Cổng", "Giai đoạn", "Chỉ số gác cổng duy nhất", "Ngưỡng định lượng bắt buộc", "Bằng chứng đối soát vật lý", "Quyết định"]
    for j, h in enumerate(headers_t3):
        cell = tbl3.cell(0, j)
        cell.width = col_widths_t3[j]
        add_styled_cell(cell, [(h, True, False, 6.8, RGBColor(255, 255, 255))], fill_hex="0F294A", top_pad=40, bot_pad=40)

    data_t3 = [
        ("Ngày 30", "Learning", "Xác nhận Pain Moment & độ chính xác RAG syllabus từ 3 Trưởng bộ môn", "≥85% Factual Accuracy & 3/3 bộ môn nghiệm thu", "Biên bản nghiệm thu KT & eval log 30 slices", "Đạt: GO\nTrượt: FIX"),
        ("Ngày 60", "Operation", "Tỷ lệ sinh viên active giải bài tập tuần (Weekly active resolution rate)", "≥50% trên tổng số 600 sinh viên tại 2 khoa pilot", "Báo cáo trích xuất telemetry log Canvas LTI", "Đạt: GO\nTrượt: PIVOT"),
        ("Ngày 90", "Business", "Gross Margin sau AI cost và chuyển đổi hợp đồng B2B chính thức", "Gross Margin ≥80% & có ≥2 hợp đồng B2B ($14.4K)", "Hợp đồng kinh tế đã ký & sao kê thanh toán", "Đạt: GO\nTrượt: KILL")
    ]

    for i, row in enumerate(data_t3, start=1):
        bg = "FFFFFF" if i % 2 != 0 else "F8FAFC"
        for j in range(6):
            tbl3.cell(i, j).width = col_widths_t3[j]
            b_flag = True if j in [0, 5] else False
            add_styled_cell(tbl3.cell(i, j), [(row[j], b_flag, False, 6.5, RGBColor(31, 41, 55))], fill_hex=bg)

    # Footer note on Page 1
    p_f = doc.add_paragraph()
    p_f.paragraph_format.space_before = Pt(3)
    p_f.paragraph_format.space_after = Pt(0)
    p_f.paragraph_format.line_spacing = 1.15
    f_runs = [
        ("Kill criteria: ", True, False, 7.0, RGBColor(185, 28, 28)),
        ("KILL dự án vào ngày 90 nếu sau 2 chu kỳ tối ưu kỹ thuật RAG mà tỷ lệ chuyển đổi POC sang hợp đồng trả phí vẫn <30% và không có khoa nào đồng ý mức giá nền tối thiểu $1.50/sinh viên/tháng.\n", False, False, 7.0, RGBColor(31, 41, 55)),
        ("Điểm chưa đo được: ", True, False, 7.0, RGBColor(15, 41, 74)),
        ("Tỷ lệ TA tiết kiệm thực tế 70% thời gian  |  Cần bảng đối soát giờ trực của TA tại 2 khoa pilot  |  Owner: Customer Success Lead  |  Ngày có số: 2026-09-30.", False, False, 7.0, RGBColor(31, 41, 55))
    ]
    for text, b, it, sz, clr in f_runs:
        r = p_f.add_run(text)
        r.font.name = "Arial"
        r.font.size = Pt(sz)
        r.font.bold = b
        r.font.italic = it
        r.font.color.rgb = clr

    # ------------------ PAGE 2: APPENDIX ------------------
    doc.add_page_break()

    p_app_title = doc.add_paragraph()
    p_app_title.paragraph_format.space_before = Pt(0)
    p_app_title.paragraph_format.space_after = Pt(1)
    r = p_app_title.add_run("PHỤ LỤC: ĐỐI SOÁT MÔ HÌNH [MH], KIỂM KÊ 11 ĐÈN & AI CRITIQUE")
    r.font.name = "Arial"
    r.font.size = Pt(11.5)
    r.font.bold = True
    r.font.color.rgb = RGBColor(15, 41, 74)

    p_app_sub = doc.add_paragraph()
    p_app_sub.paragraph_format.space_before = Pt(0)
    p_app_sub.paragraph_format.space_after = Pt(4)
    r = p_app_sub.add_run("Thuyết minh cơ sở toán học Unit Economics Day 24-25, Nguồn Benchmark & Biên bản phản biện chất lượng")
    r.font.name = "Arial"
    r.font.size = Pt(7.5)
    r.font.italic = True
    r.font.color.rgb = RGBColor(75, 85, 99)

    # Appendix 1: Model Derived Thresholds
    p_ah1 = doc.add_paragraph()
    p_ah1.paragraph_format.space_before = Pt(0)
    p_ah1.paragraph_format.space_after = Pt(2)
    r = p_ah1.add_run("A. PHỤ LỤC NGƯỠNG SUY TỪ MÔ HÌNH UNIT ECONOMICS [MH]")
    r.font.name = "Arial"
    r.font.size = Pt(8.5)
    r.font.bold = True
    r.font.color.rgb = RGBColor(15, 41, 74)

    col_widths_app1 = [Inches(0.50), Inches(1.55), Inches(1.95), Inches(1.70), Inches(1.77)]
    tbl_app1 = doc.add_table(rows=3, cols=5)
    tbl_app1.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl_app1.autofit = False
    set_table_borders(tbl_app1, color="CBD5E1", sz="4")

    headers_app1 = ["Mã ID", "Chỉ số áp dụng", "Dữ liệu đầu vào Day 24–25", "Phép tính toán học (Có số & dấu =)", "Kết quả & Ngưỡng áp dụng"]
    for j, h in enumerate(headers_app1):
        cell = tbl_app1.cell(0, j)
        cell.width = col_widths_app1[j]
        add_styled_cell(cell, [(h, True, False, 6.8, RGBColor(255, 255, 255))], fill_hex="0F294A", top_pad=40, bot_pad=40)

    data_app1 = [
        ("MH-01", "Chi phí AI tối đa / resolution\n(Áp dụng O-02 & G-01)", "Giá bán đề xuất: 9.100 đ/job ($0.35)\nGross margin mục tiêu: 84%\nChi phí biến đổi khác: 600 đ/job", "9.100 × (1 − 0.84) − 600 = 856 đ\nTrần chịu đựng: 9.100 × (1 − 0.65) − 600 = 2.585 đ", "Xanh: ≤ 1.000 đ/job (GM >82%)\nVàng: 1.001–1.800 đ (GM 70–80%)\nĐỏ: > 1.800 đ (Vi phạm trần GM)"),
        ("MH-02", "Tỷ lệ kích hoạt Pilot tối thiểu\n(Áp dụng O-01)", "Mục tiêu doanh thu: Ký 3 hợp đồng B2B ($14.400 ACV) từ phễu 4 pilot khoa", "3 ÷ 4 = 75%\nMức hòa vốn GTM: 2 ÷ 4 = 50%", "Xanh: ≥ 75% (Đạt mục tiêu 3 deal)\nVàng: 50–74% (Cần tối ưu sales)\nĐỏ: < 50% (Không thể đạt ARR)")
    ]

    for i, row in enumerate(data_app1, start=1):
        bg = "FFFFFF" if i % 2 != 0 else "F8FAFC"
        for j in range(5):
            tbl_app1.cell(i, j).width = col_widths_app1[j]
            b_flag = True if j in [0, 1] else False
            add_styled_cell(tbl_app1.cell(i, j), [(row[j], b_flag, False, 6.5, RGBColor(31, 41, 55))], fill_hex=bg)

    # Appendix 2: 11 Candidates
    p_ah2 = doc.add_paragraph()
    p_ah2.paragraph_format.space_before = Pt(4)
    p_ah2.paragraph_format.space_after = Pt(2)
    r = p_ah2.add_run("B. KIỂM KÊ TOÀN BỘ 11 ĐÈN ỨNG VIÊN B2B TỪ HANDBOOK")
    r.font.name = "Arial"
    r.font.size = Pt(8.5)
    r.font.bold = True
    r.font.color.rgb = RGBColor(15, 41, 74)

    col_widths_app2 = [Inches(1.80), Inches(0.40), Inches(0.70), Inches(4.57)]
    tbl_app2 = doc.add_table(rows=12, cols=4)
    tbl_app2.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl_app2.autofit = False
    set_table_borders(tbl_app2, color="CBD5E1", sz="4")

    headers_app2 = ["Đèn ứng viên B2B từ Handbook", "Tầng", "Trạng thái", "Bằng chứng hiện có / Kế hoạch đo lường chi tiết"]
    for j, h in enumerate(headers_app2):
        cell = tbl_app2.cell(0, j)
        cell.width = col_widths_app2[j]
        add_styled_cell(cell, [(h, True, False, 6.8, RGBColor(255, 255, 255))], fill_hex="0F294A", top_pad=35, bot_pad=35)

    data_app2 = [
        ("Time-to-first-value (TTFV)", "L", "Đo được", "Log tích hợp LTI 1.3 và event 50 resolution đầu tiên của sinh viên"),
        ("Pipeline coverage", "L", "Đang đo", "Chuẩn hóa stage cơ hội và CRM pipeline của 8 khoa mục tiêu trước 2026-09-15"),
        ("% deal chết ở khâu security/procurement", "L", "Đang đo", "Theo dõi lý do từ chối về DPA/bảo mật RAG trong CRM trước 2026-09-15"),
        ("POC → paid", "O", "Đo được", "Biên bản MOU và kết quả chuyển đổi từ 3 pilot sang hợp đồng chính thức"),
        ("Sales cycle (ngày)", "O", "Đang đo", "Đo số tuần từ demo LTI đến ký hợp đồng mua sắm khoa trước 2026-09-30"),
        ("Usage depth trong tài khoản", "O", "Đo được", "Telemetry log tỷ lệ sinh viên kích hoạt giải bài tập tuần trên Canvas"),
        ("Chi phí triển khai ÷ ACV", "O", "Đang đo", "Timesheet hỗ trợ cấu hình syllabus và RAG index trước 2026-09-20"),
        ("Tập trung doanh thu", "O", "Đo được", "Bảng theo dõi ACV 3 hợp đồng khoa đầu tiên so với tổng ARR"),
        ("NRR", "G", "Đang đo", "Đo lường mở rộng số lớp/môn sau 2 học kỳ trước 2027-03-31"),
        ("Gross Margin", "G", "Đo được", "Đối soát billing token Claude Haiku và chi phí hạ tầng vector DB"),
        ("CAC payback", "G", "Đang đo", "Tính fully-loaded CAC sau khi hoàn thành 3 thương vụ đầu trước 2026-10-31")
    ]

    for i, row in enumerate(data_app2, start=1):
        bg = "FFFFFF" if i % 2 != 0 else "F8FAFC"
        for j in range(4):
            tbl_app2.cell(i, j).width = col_widths_app2[j]
            b_flag = True if j in [0, 2] else False
            color_c = RGBColor(4, 120, 87) if row[j] == "Đo được" else RGBColor(31, 41, 55)
            add_styled_cell(tbl_app2.cell(i, j), [(row[j], b_flag, False, 6.4, color_c)], fill_hex=bg)

    # Appendix 3: AI Critique
    p_ah3 = doc.add_paragraph()
    p_ah3.paragraph_format.space_before = Pt(4)
    p_ah3.paragraph_format.space_after = Pt(2)
    r = p_ah3.add_run("C. GHI NHẬN AI CRITIQUE & PHẢN BIỆN CHUYÊN MÔN")
    r.font.name = "Arial"
    r.font.size = Pt(8.5)
    r.font.bold = True
    r.font.color.rgb = RGBColor(15, 41, 74)

    col_widths_app3 = [Inches(2.10), Inches(0.85), Inches(2.20), Inches(2.32)]
    tbl_app3 = doc.add_table(rows=4, cols=4)
    tbl_app3.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl_app3.autofit = False
    set_table_borders(tbl_app3, color="CBD5E1", sz="4")

    headers_app3 = ["Ý kiến phản biện AI", "Quyết định", "Thay đổi đã thực hiện", "Lý do & Thuyết minh chuyên môn"]
    for j, h in enumerate(headers_app3):
        cell = tbl_app3.cell(0, j)
        cell.width = col_widths_app3[j]
        add_styled_cell(cell, [(h, True, False, 6.8, RGBColor(255, 255, 255))], fill_hex="0F294A", top_pad=40, bot_pad=40)

    data_app3 = [
        ("TTFV cho trường đại học cần đo theo số resolution học tập thay vì chỉ số lượng SV đăng nhập", "Chấp nhận", "Định nghĩa TTFV là thời gian đạt 50 resolution học tập đạt chuẩn QA", "Tránh bẫy vanity metric khi sinh viên đăng nhập nhưng không thực sự học tập"),
        ("Nên lấy benchmark NRR SaaS chung 101% làm ngưỡng đỏ", "Bác bỏ", "Giữ ngưỡng đỏ là <100% và xanh là ≥110%", "Trong B2B giáo dục, nếu NRR <100% nghĩa là khoa bị cắt giảm môn học đăng ký"),
        ("Cần có cơ chế kiểm soát token cost tránh sinh viên spam câu hỏi ngoài bài học", "Chấp nhận", "Bổ sung rule R-04 kích hoạt prompt caching & dynamic routing", "Bảo vệ biên lợi nhuận gộp 84% khỏi nguy cơ đội chi phí inference")
    ]

    for i, row in enumerate(data_app3, start=1):
        bg = "FFFFFF" if i % 2 != 0 else "F8FAFC"
        for j in range(4):
            tbl_app3.cell(i, j).width = col_widths_app3[j]
            b_flag = True if j == 1 else False
            color_c = RGBColor(4, 120, 87) if row[j] == "Chấp nhận" else (RGBColor(185, 28, 28) if row[j] == "Bác bỏ" else RGBColor(31, 41, 55))
            add_styled_cell(tbl_app3.cell(i, j), [(row[j], b_flag, False, 6.4, color_c)], fill_hex=bg)

    docx_path = "TrinhHaiDang_Day26_dashboard.docx"
    doc.save(docx_path)
    print(f"Saved {docx_path} successfully.")

if __name__ == '__main__':
    create_document()
